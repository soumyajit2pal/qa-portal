from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "docs" / "QA_Portal_User_Manual.docx"
NAVY = RGBColor(7, 75, 91)
TEAL = RGBColor(13, 102, 120)
GOLD = RGBColor(196, 151, 28)
MUTED = RGBColor(95, 108, 113)
LIGHT = "E8F1F2"


def font(run, size=11, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def configure_table(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_pr.append(tbl_layout)
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), "9360")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            set_cell_margins(cell)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, text in enumerate(headers):
        set_cell_fill(table.rows[0].cells[i], LIGHT)
        p = table.rows[0].cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        font(p.add_run(text), 9.5, True, NAVY)
    repeat_header(table.rows[0])
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            font(p.add_run(str(value)), 9.5)
    configure_table(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        font(p.add_run(item), 10.5)


def add_steps(doc, items):
    numbering = doc.part.numbering_part.element
    base_num_id = doc.styles["List Number"]._element.pPr.numPr.numId.val
    base_num = next(node for node in numbering.findall(qn("w:num")) if node.get(qn("w:numId")) == str(base_num_id))
    abstract_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    next_num_id = max(int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))) + 1
    num = OxmlElement("w:num"); num.set(qn("w:numId"), str(next_num_id))
    abstract = OxmlElement("w:abstractNumId"); abstract.set(qn("w:val"), abstract_id); num.append(abstract)
    override = OxmlElement("w:lvlOverride"); override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride"); start.set(qn("w:val"), "1"); override.append(start); num.append(override)
    numbering.append(num)
    for title, detail in items:
        p = doc.add_paragraph(style="List Number")
        num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
        num_pr.get_or_add_ilvl().val = 0
        num_pr.get_or_add_numId().val = next_num_id
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        font(p.add_run(f"{title}. "), 10.5, True, NAVY)
        font(p.add_run(detail), 10.5)


def add_note(doc, title, text, warning=False):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_fill(cell, "FFF6DA" if warning else "EDF6F7")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    font(p.add_run(title + "\n"), 10.5, True, GOLD if warning else TEAL)
    font(p.add_run(text), 10)
    configure_table(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


doc = Document()
section = doc.sections[0]
section.page_width, section.page_height = Inches(8.5), Inches(11)
section.top_margin = section.bottom_margin = Inches(1)
section.left_margin = section.right_margin = Inches(1)
section.header_distance = section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.25
for name, size, before, after, color in (
    ("Heading 1", 16, 18, 10, NAVY), ("Heading 2", 13, 14, 7, TEAL), ("Heading 3", 12, 10, 5, NAVY)
):
    style = styles[name]
    style.font.name = "Calibri"; style.font.size = Pt(size); style.font.bold = True; style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
font(header.add_run("QualityHub | Bank of Maharashtra QA Portal"), 8.5, True, MUTED)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(footer.add_run("Controlled User Guide  |  "), 8, False, MUTED)
fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
footer._p.append(fld)

# Editorial cover
doc.add_paragraph().paragraph_format.space_after = Pt(80)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("OPERATING GUIDE"), 11, True, GOLD)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(8)
font(p.add_run("QualityHub"), 30, True, NAVY)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(8)
font(p.add_run("QA Portal User Manual"), 18, False, TEAL)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(65)
font(p.add_run("Requests, approvals, evidence, test management, execution, reporting, and administration"), 11, False, MUTED, True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("Bank of Maharashtra\nQuality Assurance Department - IT\nVersion 1.0 | 07 August 2026"), 11, True, NAVY)
doc.add_page_break()

heading(doc, "How to use this manual", 1)
doc.add_paragraph("This manual is a role-aware operating guide for the current QualityHub portal. Screens and actions can differ by role, department, assignment, project status, and workflow stage.")
add_note(doc, "Fastest support path", "Capture the full TQA record ID, module, current status, Pending With value, exact popup reason, timestamp, and a masked screenshot. Never send passwords or secrets.")
heading(doc, "Contents", 2)
add_table(doc, ["Section", "Purpose"], [
    ("1. Getting started", "Navigation, dashboard, drawers, tables, and search"),
    ("2. Roles and access", "Role responsibilities, department scope, and controls"),
    ("3. QA request intake", "Draft, submit, evidence, linked child requests"),
    ("4. Approval workflows", "Approve, return, reject, lifecycle branches"),
    ("5. Test management", "Projects, repository, tags, import, bulk maintenance"),
    ("6. Test execution", "Cycles, request linking, runners, attempts, defects"),
    ("7. Evidence and collaboration", "Documents, rich text, images, activity"),
    ("8. Monitoring and reporting", "Dashboard, Insights, approvals, exports"),
    ("9. Administration", "Users, departments, checklists, upload storage"),
    ("10. Troubleshooting", "Common symptoms and corrective actions"),
], [2500, 6860])

heading(doc, "1. Getting started", 1)
add_steps(doc, [
    ("Sign in", "Use the Standard or LDAP/Active Directory method assigned to the account"),
    ("Verify profile", "Open the user menu and confirm department and every assigned role"),
    ("Review work", "Open Dashboard and Pending Approvals before starting a new request"),
    ("Use full IDs", "Use complete TQA IDs in searches, comments, exports, and support requests"),
])
heading(doc, "Navigation and work surfaces", 2)
add_bullets(doc, [
    "The left navigation is grouped and collapsible; active modules are highlighted.",
    "Request drawers open expanded and remain open until Close is selected. Use the expand/unexpand control to change width.",
    "Global search routes known TQA prefixes directly to their correct module and record.",
    "Dashboard shows operational attention items; Insights contains focused Security, Suppression, and 3W analytics.",
])
heading(doc, "Tables", 2)
add_bullets(doc, [
    "Designed columns are visible by default. Open Columns to add other fields returned by the API or hide visible fields.",
    "Column choices are saved per table for the current browser. At least one column must remain visible.",
    "Use the filter icon in a column heading for field-specific filtering. Tables paginate automatically.",
    "The Columns popup scrolls independently and stays within the viewport.",
])

heading(doc, "2. Roles and access", 1)
add_table(doc, ["Role", "Primary responsibility", "Scope/control"], [
    ("Requester / Business Analyst", "Raise requests, supply details/evidence, correct returns, verify completion", "Own or department work"),
    ("Application Owner", "Approve/reject proposed application names", "Same department"),
    ("SM", "First business review", "Same department; no self-approval"),
    ("Department Head", "Business approval and QA Lead assignment", "Same department"),
    ("QA Engineer", "Author test cases and execute assigned testing", "IT - QA"),
    ("QA Lead", "Readiness, assignment, test-case approval, delivery governance", "IT - QA cross-department"),
    ("Security Analyst", "SAST/DAST execution, finding validation, suppression review", "IT - QA"),
    ("Executive COE", "Final QA Sign-off and QA access coordination", "IT - QA governance"),
    ("Administrator", "Users, departments, settings, checklists, protected access", "System-wide"),
], [1700, 4760, 2900])
add_note(doc, "Multiple roles", "Roles are additive; users do not switch personas. Stage, department, assignment, account state, and segregation-of-duties rules still apply.", True)

heading(doc, "3. QA request intake", 1)
add_steps(doc, [
    ("Prepare", "Collect change reference, application, environment, release, contacts, scope, risk, and evidence"),
    ("Select testing", "Choose Functional, SAST, DAST, and/or Performance"),
    ("Complete declarations", "Answer readiness criteria and attach supporting evidence"),
    ("Save or submit", "Draft remains editable; Submit validates and raises linked child requests"),
    ("Track children", "Operational progress occurs on TQA-FUNC, TQA-SAST, TQA-DAST, and TQA-PERF records"),
])
add_note(doc, "Folder behavior", "Draft uploads remain traceable when the generated TQA request ID is assigned. Storage paths are normalized to prevent repeated or duplicate folder structures.")

heading(doc, "4. Approval and lifecycle workflows", 1)
add_table(doc, ["Decision", "Meaning", "Expected user action"], [
    ("Approve", "Stage is acceptable and advances", "Review details and optionally add a clear action note"),
    ("Return", "Correction is required", "State exactly what must change; requester edits and resubmits"),
    ("Reject", "Request should not proceed", "Give the business, control, or technical reason"),
], [1500, 3100, 4760])
heading(doc, "Functional lifecycle", 2)
doc.add_paragraph("Requester -> SM -> Department Head -> QA Lead readiness -> QA Tester -> QA Sign-off -> Requester verification -> Closed")
add_bullets(doc, [
    "Returned requests show Requester Action as the current lifecycle destination, not the reviewer’s previous stage.",
    "Rejected, cancelled, or early-closed requests connect directly to Closed; unreached QA and Sign-off stages are not shown as complete.",
    "Pending Approvals groups child requests beneath their parent request ID and opens the selected item directly.",
])
heading(doc, "Other workflows", 2)
add_bullets(doc, [
    "SAST/DAST: SM -> Department Head -> QA readiness -> Security Analyst -> findings/remediation/rescan -> Report Ready.",
    "Performance: SM -> Department Head -> readiness -> planning/execution -> analysis/report -> sign-off -> Closed.",
    "Suppression: SM -> Department Head -> Security verification -> Done or Rejected.",
    "QA Sign-off: QA Engineer -> QA Lead -> Executive COE -> Issued; no SM stage.",
])

heading(doc, "5. Test management", 1)
heading(doc, "Projects and repository", 2)
add_bullets(doc, [
    "Department is mandatory and selected from the system department list.",
    "When an Application is selected, its mapped Department is applied and the Department field becomes read-only.",
    "Folders and subfolders organize release, module, epic, or scope. Tags/labels support testcase filtering.",
    "Bulk update can change Test Type, Folder, Module Name, and Priority.",
])
heading(doc, "Excel import", 2)
add_steps(doc, [
    ("Download template", "Use the standard Test Case Import Template.xlsx"),
    ("Choose destination", "Select the active project and optional target folder"),
    ("Upload and validate", "The workbook is parsed and each testcase becomes Draft/Pending QA Lead Review"),
    ("Review result", "Correct skipped rows using the exact displayed reason before retrying"),
])
add_note(doc, "Import-file retention", "The imported Excel workbook is read in memory and is not stored under the upload path. Test cases and steps are saved in the database.")
heading(doc, "Review and version control", 2)
add_bullets(doc, [
    "QA Lead approval makes a testcase Active and eligible for cycles.",
    "Check Out reserves editing; Check In releases the reservation.",
    "Materially updated testcases return to review and maintain governed version history.",
])

heading(doc, "6. Test cycles and execution", 1)
add_steps(doc, [
    ("Create or edit cycle", "Set scope and dates, and optionally link a child Functional, SAST, DAST, or Performance request ID"),
    ("Add testcases", "Use Select all, filters, pagination, and configurable columns; only Active unlinked cases are eligible"),
    ("Assign runners", "Assign or reassign active IT-QA QA Engineers or QA Leads"),
    ("Execute", "Use quick Play or the detailed runner to record a retained attempt"),
    ("Record result", "Capture status, actual result, notes, and supported evidence"),
    ("Link defect", "Fail and Blocked outcomes can link a defect only to the latest applicable failed/blocked run"),
])
heading(doc, "Request and lifecycle links", 2)
add_bullets(doc, [
    "Starting Functional execution asks whether an eligible test cycle should be linked. Cancel makes no status change.",
    "A linked cycle is visible on request details and duplicate links are prevented.",
    "Unlink can be performed from Functional Request details or Test Lifecycle without deleting testcases or run history.",
    "Reset Lifecycle requires a destructive confirmation and returns cycle execution state to its initial condition.",
])

heading(doc, "7. Evidence, rich text, and collaboration", 1)
add_bullets(doc, [
    "Documents and checklist evidence use the configured upload root and are stored against governed record folders.",
    "Execution attempts can retain screenshots and evidence with the individual run.",
    "Activity supports formatted comments, lists, pasted images, and attachments; workflow decisions retain their audit note.",
    "QA Certificate Exit Criteria Validation Notes, Open Defect Review Summary, and Residual Risk Documentation support rich text and tables.",
    "QA Certificate PDF export preserves formatted text and table data.",
])
add_note(doc, "Information security", "Do not upload credentials, secrets, production customer data, or unmasked personal information.", True)

heading(doc, "8. Monitoring and reporting", 1)
add_bullets(doc, [
    "Dashboard prioritizes active projects, findings, decisions, active child requests, governance risks, lifecycle health, and recent activity.",
    "Insights consolidates Security, Suppression, and 3W Pending analytics.",
    "Pending Approvals refreshes after a decision and groups children under parent requests.",
    "Reports and exports provide governed summaries; Repository and Test Execution export detailed assets and results.",
    "TQA-REQ gateway records are not double-counted as executable child requests in dashboard request totals.",
])

heading(doc, "9. Administration", 1)
heading(doc, "Users and departments", 2)
add_bullets(doc, [
    "Create, activate/deactivate, and review accounts without deleting audit history.",
    "Set the correct Department before assigning roles; periodically certify multi-role and privileged access.",
    "Reassign open requests and testing work before deactivation or transfer.",
])
heading(doc, "Upload storage", 2)
add_steps(doc, [
    ("Open System Settings", "Enter an absolute server filesystem path"),
    ("Validate access", "The API verifies the path exists or can be created and is writable"),
    ("Mount persistently", "In production, mount the same persistent volume at that exact path in every API replica"),
    ("Retain legacy roots", "Keep configured legacy paths available while older documents still reference them"),
])
add_note(doc, "Default development path", "If no Admin setting exists, retained uploads use backend/app/uploads. Production should use a durable mounted path outside the application image.", True)

heading(doc, "10. Troubleshooting", 1)
add_table(doc, ["Symptom", "Likely reason", "Corrective action"], [
    ("Missing action", "Wrong stage, role, department, assignment, or inactive account", "Check profile, status, Pending With, and assignment"),
    ("400 popup", "Required data or workflow precondition missing", "Follow the exact reason and guidance"),
    ("403 popup", "Authenticated but not authorized", "Verify role, department, self-approval, assignment"),
    ("404", "Wrong route/ID or missing record", "Search full TQA ID in the correct module"),
    ("Import skipped", "Missing/duplicate/invalid template data", "Correct rows using the displayed issue summary"),
    ("Cannot execute", "Unapproved case, inactive project, or no runner", "Approve/reactivate/assign as authorized"),
    ("Upload path error", "Path not absolute, writable, or mounted", "Correct Admin path and container volume mapping"),
    ("PDF formatting issue", "Unsupported or malformed rich content", "Use supported editor formatting and regenerate"),
], [2200, 3350, 3810])
heading(doc, "Escalation checklist", 2)
add_bullets(doc, [
    "Full TQA record ID, module name, current status, and Pending With value.",
    "Department and roles (never a password), plus the exact popup reason and corrective guidance.",
    "Timestamp, attempted action, and a masked screenshot where useful.",
])

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.core_properties.title = "QualityHub QA Portal User Manual"
doc.core_properties.subject = "Current user and administrator operating guide"
doc.core_properties.author = "Quality Assurance Department - IT"
doc.core_properties.keywords = "QualityHub, QA Portal, user manual, test management, approvals"
doc.save(OUT)
print(OUT)
